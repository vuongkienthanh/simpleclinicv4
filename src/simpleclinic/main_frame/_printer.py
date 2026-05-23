"""
READ fill_main_frame_data for available key:value
"""

from docx import Document, document
import copy
import wx
import re
from python_docx_replace import docx_replace
from lib.paths import PRESCRIPTION_TEMPLATE, PRESCRIPTION_OUTPUT
from lib.vn import wxdatetime_to_vietnamese_date
from lib.wx_helper import get_main_frame
from lib import DATE_FORMAT


def _duplicate_medicine_rows(doc: document.Document, count: int):
    # Find the index of the "Thuốc:" paragraph
    thuoc_index = -1
    for i, p in enumerate(doc.paragraphs):
        if "Thuốc:" in p.text:
            thuoc_index = i
            break

    if thuoc_index == -1:
        print("Could not find the 'Thuốc:' paragraph in the document.")
        return

    # Grab the target templates
    try:
        row1_template = doc.paragraphs[thuoc_index + 1]
        row2_template = doc.paragraphs[thuoc_index + 2]
    except IndexError:
        print("The document does not have 2 rows after 'Thuốc:'.")
        return

    placeholder_pattern = re.compile(r"\${([^}]+)}")
    current_element = doc.paragraphs[thuoc_index]._p

    # Loop backward to naturally stack elements downward using .addnext()
    for index in reversed(range(count)):
        # Inline function to handle index string assignment (no underscore)
        def add_index(match):
            key = match.group(1)
            return f"${{{key}{index}}}"

        # --- Process and Clone Row 2 ---
        p2_element = copy.deepcopy(row2_template._element)
        current_element.addnext(p2_element)
        # Re-wrap xml element into a python-docx Paragraph object
        for run in p2_element.xpath(".//w:r"):
            text_elements = run.xpath(".//w:t")
            for t in text_elements:
                if t.text:
                    t.text = placeholder_pattern.sub(add_index, t.text)

        # --- Process and Clone Row 1 ---
        p1_element = copy.deepcopy(row1_template._element)
        current_element.addnext(p1_element)
        for run in p1_element.xpath(".//w:r"):
            text_elements = run.xpath(".//w:t")
            for t in text_elements:
                if t.text:
                    t.text = placeholder_pattern.sub(add_index, t.text)

    # --- DELETE ORIGINAL TEMPLATE ROWS ---
    r1_p = row1_template._element
    r2_p = row2_template._element
    r1_p.getparent().remove(r1_p)
    r2_p.getparent().remove(r2_p)


def fill_main_frame_data():
    d = Document(str(PRESCRIPTION_TEMPLATE))
    _duplicate_medicine_rows(d, get_main_frame().medicine_list.ItemCount)
    date = wx.DateTime()
    date.ParseISOCombined(get_main_frame().visit_exam_datetime.Value)
    data = {
        "name": get_main_frame().patient_name.Value,
        "gender": get_main_frame().patient_gender.GetGender().display_name,
        "birthdate": get_main_frame().patient_birthdate.Value.Format(DATE_FORMAT),
        "weight": str(get_main_frame().visit_weight.GetInt()),
        "diagnosis": get_main_frame().visit_diagnosis.Value.strip(),
        "days": str(get_main_frame().visit_days.Value),
        "note": get_main_frame().visit_note.Value.strip(),
        "date": wxdatetime_to_vietnamese_date(date),
    }
    for i in range(get_main_frame().medicine_list.ItemCount):
        data |= {
            f"STT{i}": str(i + 1),
            f"mname{i}": get_main_frame().medicine_list.GetItemText(i, 2),
            f"element{i}": get_main_frame().medicine_list.GetItemText(i, 3),
            f"quantity{i}": get_main_frame().medicine_list.GetItemText(i, 6),
            f"usage{i}": "{} ngày {}, lần {} ({})".format(
                get_main_frame().medicine_list.GetItemText(i, 4),
                get_main_frame().medicine_list.GetItemText(i, 5),
                get_main_frame().medicine_list.GetItemText(i, 6),
                get_main_frame().medicine_list.GetItemText(i, 8),
            ),
        }
    docx_replace(data)
    d.save(str(PRESCRIPTION_OUTPUT))
